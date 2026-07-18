#!/usr/bin/env python3
"""THROWAWAY prototype: canonical ReportView -> Markdown -> validated DOCX."""
from __future__ import annotations

import re
import shutil
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Literal
from urllib.parse import urlparse

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
TMP = ROOT / ".tmp-run"
CONCLUSIONS = ("normal", "warning", "abnormal", "undetermined", "not_applicable")
LABELS = {"normal": "正常", "warning": "警告", "abnormal": "异常", "undetermined": "无法判断", "not_applicable": "不适用"}
DOMAINS = ("主机资源", "Tomcat 实例与 JVM", "Connector 与线程池", "Tomcat 配置安全", "应用部署", "日志配置与文件状态")
MAJOR_SECTIONS = ("报告说明", "实例概况", "结论摘要", "六巡检域", "无法判断与采集限制", "附录")
METADATA = {"protocol_version": "1.0", "script_version": "prototype-0.2", "ruleset_id": "tomcat_mvp", "ruleset_version": "1.0.0"}


@dataclass(frozen=True)
class InstanceIdentity:
    instance_id: str
    hostname: str
    tomcat_version: str
    install_path_hint: str
    process_user: str
    http_port: str


@dataclass(frozen=True)
class RuleResult:
    rule_id: str
    domain: str
    subject: str
    conclusion: Literal["normal", "warning", "abnormal", "undetermined", "not_applicable"]
    reason_codes: tuple[str, ...]
    fact_refs: tuple[str, ...]
    related_subjects: tuple[str, ...] = ()
    evidence_limitations: tuple[str, ...] = ()
    recommendation_ids: tuple[str, ...] = ()
    verification_ids: tuple[str, ...] = ()

    def validate(self) -> None:
        if self.conclusion not in CONCLUSIONS or not self.reason_codes or not self.fact_refs:
            raise ValueError("rule result lacks canonical conclusion/reason/fact")
        if self.conclusion in {"warning", "abnormal"}:
            if not self.recommendation_ids or self.verification_ids:
                raise ValueError("warning/abnormal require recommendation only")
        elif self.conclusion == "undetermined":
            if not self.verification_ids or self.recommendation_ids:
                raise ValueError("undetermined requires verification only")
        elif self.recommendation_ids or self.verification_ids:
            raise ValueError("normal/not_applicable cannot carry recommendation or verification")


@dataclass(frozen=True)
class Observation:
    observation_id: str
    domain: str
    text: str
    fact_ref: str


@dataclass(frozen=True)
class SharedResultReference:
    rule_id: str
    subject: str
    conclusion: str
    reason_code: str


@dataclass(frozen=True)
class ReportView:
    protocol_version: str
    script_version: str
    collected_at: str
    ruleset_id: str
    ruleset_version: str
    instance: InstanceIdentity
    user_revised: bool  # Always false: this is the immutable initial system view.
    summary_counts: dict[str, int]
    shared_result_references: tuple[SharedResultReference, ...]
    rules: tuple[RuleResult, ...]
    observations: tuple[Observation, ...]
    limitations: tuple[str, ...]
    appendix: tuple[str, ...]

    def validate(self) -> None:
        for key, expected in METADATA.items():
            if getattr(self, key) != expected:
                raise ValueError(f"metadata {key} mismatch")
        if self.user_revised:
            raise ValueError("ReportView is the initial system view; user_revised must be false")
        if self.instance.instance_id == "aggregate" or set(self.summary_counts) != set(CONCLUSIONS):
            raise ValueError("invalid report scope or summary")
        seen: set[tuple[str, str]] = set()
        counts = {code: 0 for code in CONCLUSIONS}
        for rule in self.rules:
            rule.validate()
            if rule.domain not in DOMAINS or (rule.rule_id, rule.subject) in seen:
                raise ValueError("invalid domain or duplicate rule subject")
            seen.add((rule.rule_id, rule.subject))
            counts[rule.conclusion] += 1
        if counts != self.summary_counts:
            raise ValueError("summary must count rules only")


def rule(rule_id: str, domain: str, subject: str, conclusion: str, reason: str, fact: str, *, rec: str = "", verify: str = "", limitation: str = "") -> RuleResult:
    return RuleResult(rule_id, domain, subject, conclusion, (reason,), (fact,), (), (limitation,) if limitation else (), (rec,) if rec else (), (verify,) if verify else ())


def make_report(instance_id: str, version: str, degraded: bool = False) -> ReportView:
    base = f"/instances/{0 if instance_id == 'A' else 1}"
    connector = f"{base}/connectors/0"
    app = f"{base}/applications/0"
    if degraded:
        rules = (
            rule("HOST.FILESYSTEM_CAPACITY", "主机资源", f"{base}/filesystems/0", "abnormal", "filesystem_resource_exhausted", f"{base}/filesystems/0/available_bytes", rec="restore_exhausted_host_resource"),
            rule("SECURITY.VERSION_LIFECYCLE", "Tomcat 实例与 JVM", base, "normal", "tomcat_branch_supported", f"{base}/identity/tomcat_version"),
            rule("JVM.VERSION_COMPATIBILITY", "Tomcat 实例与 JVM", base, "undetermined", "compatibility_evidence_incomplete", f"{base}/jvm/java_feature_version", verify="verify_exact_tomcat_and_java_versions", limitation="java_version_unavailable"),
            rule("CONNECTOR.TIMEOUTS", "Connector 与线程池", connector, "abnormal", "connector_timeout_invalid", f"{connector}/connection_timeout", rec="correct_invalid_effective_configuration"),
            rule("SECURITY.SHUTDOWN_PORT", "Tomcat 配置安全", base, "warning", "shutdown_port_enabled", f"{base}/server/shutdown_port", rec="disable_or_harden_shutdown_port"),
            rule("SECURITY.AJP", "Tomcat 配置安全", base, "not_applicable", "no_ajp_connector", f"{base}/connectors"),
            rule("APPLICATION.CONTEXT_CONFLICT", "应用部署", app, "abnormal", "application_context_conflict", f"{app}/context_path", rec="resolve_application_context_conflict"),
            rule("LOGGING.ROTATION", "日志配置与文件状态", f"{base}/logging/targets/0", "undetermined", "logging_rotation_unresolved", f"{base}/logging/targets/0/rotation", verify="verify_logging_configuration_and_paths", limitation="external_rotation_visibility_unknown"),
        )
    else:
        rules = (
            rule("HOST.FILESYSTEM_CAPACITY", "主机资源", f"{base}/filesystems/0", "normal", "filesystem_resource_not_exhausted", f"{base}/filesystems/0/available_bytes"),
            rule("SECURITY.VERSION_LIFECYCLE", "Tomcat 实例与 JVM", base, "normal", "tomcat_branch_supported", f"{base}/identity/tomcat_version"),
            rule("JVM.VERSION_COMPATIBILITY", "Tomcat 实例与 JVM", base, "normal", "java_tomcat_compatible", f"{base}/jvm/java_feature_version"),
            rule("CONNECTOR.TIMEOUTS", "Connector 与线程池", connector, "normal", "connector_timeout_valid", f"{connector}/connection_timeout"),
            rule("SECURITY.SHUTDOWN_PORT", "Tomcat 配置安全", base, "normal", "shutdown_port_disabled", f"{base}/server/shutdown_port"),
            rule("SECURITY.AJP", "Tomcat 配置安全", base, "not_applicable", "no_ajp_connector", f"{base}/connectors"),
            rule("APPLICATION.CONTEXT_CONFLICT", "应用部署", app, "normal", "application_context_unique", f"{app}/context_path"),
            rule("LOGGING.ROTATION", "日志配置与文件状态", f"{base}/logging/targets/0", "normal", "logging_rotation_configured", f"{base}/logging/targets/0/rotation"),
        )
    counts = {code: sum(r.conclusion == code for r in rules) for code in CONCLUSIONS}
    host = f"tomcat-{instance_id}"
    report = ReportView(
        **METADATA,
        collected_at="2026-07-18T10:15:00+08:00" if instance_id == "A" else "2026-07-18T10:20:00+08:00",
        instance=InstanceIdentity(instance_id, host, version, f"/srv/{host}", "tomcat", "8080" if instance_id == "A" else "18080"),
        user_revised=False,
        summary_counts=counts,
        shared_result_references=(
            SharedResultReference("INSTANCE.DISCOVERY", "/discovery", "normal", "discovery_complete"),
            SharedResultReference("HOST.VISIBILITY", "/host", "normal", "host_visibility_complete"),
            SharedResultReference("HOST.MEMORY_CAPACITY", "/host", "normal", "available_memory_not_exhausted"),
        ),
        rules=rules,
        observations=(
            Observation("HOST.CPU_CAPACITY", "主机资源", "CPU 核数、瞬时 CPU 与系统负载仅作观察。", "/host/cpu"),
            Observation("JVM.MEMORY_OPTIONS", "Tomcat 实例与 JVM", "堆大小与进程线程数仅作观察。", f"{base}/jvm/memory_options"),
            Observation("CONNECTOR.LIMITS", "Connector 与线程池", "线程、连接与队列有效值仅作观察。", f"{connector}/limits"),
            Observation("APPLICATION.INVENTORY", "应用部署", "应用数量仅作观察。", f"{base}/applications"),
            Observation("LOGGING.CONFIGURATION", "日志配置与文件状态", "日志文件大小和修改时间仅作观察。", f"{base}/logging/files"),
        ),
        limitations=("本报告基于一次性配置与状态快照，不替代持续监控、性能压测或故障诊断。", "受支持分支不等于当前补丁安全。", "不包含密钥、原始配置、日志正文或原始命令。"),
        appendix=("五类结论分别计数；观察项不进入统计。", "共享结果只限发现、主机可见性和主机内存等批次结果。"),
    )
    report.validate()
    return report


def render_report_view_markdown(report: ReportView) -> str:
    report.validate()
    lines = [f"# Tomcat 巡检报告 - {report.instance.instance_id}", "", "**报告来源**: 系统生成报告", "", "## 报告说明", "", *report.limitations, "", "## 实例概况", "", "| 字段 | 值 |", "| --- | --- |"]
    metadata = [(k, getattr(report, k)) for k in ("protocol_version", "script_version", "collected_at", "ruleset_id", "ruleset_version")]
    metadata += [("instance_id", report.instance.instance_id), ("hostname", report.instance.hostname), ("tomcat_version", report.instance.tomcat_version), ("install_path_hint", report.instance.install_path_hint), ("process_user", report.instance.process_user), ("http_port", report.instance.http_port), ("user_revised", "false")]
    lines += [f"| {k} | {v} |" for k, v in metadata]
    lines += ["", "## 结论摘要", "", "| 状态 | 数量 |", "| --- | --- |"]
    lines += [f"| {LABELS[code]} | {report.summary_counts[code]} |" for code in CONCLUSIONS]
    lines += ["", "观察项为非判断性事实记录，不计入规则结果数量。", "本报告不提供总体风险、评级或评分。", "", "## 六巡检域", ""]
    for domain in DOMAINS:
        lines += [f"### {domain}", "", "| 规则ID | subject | 结论 | reason_codes | fact_refs |", "| --- | --- | --- | --- | --- |"]
        domain_rules = [r for r in report.rules if r.domain == domain]
        for r in domain_rules:
            lines.append(f"| {r.rule_id} | `{r.subject}` | {LABELS[r.conclusion]} | `{', '.join(r.reason_codes)}` | `{', '.join(r.fact_refs)}` |")
        lines.append("")
        for r in domain_rules:
            if r.evidence_limitations:
                lines.append(f"- **{r.rule_id} 证据限制**: `{', '.join(r.evidence_limitations)}`")
            if r.recommendation_ids:
                lines.append(f"- **{r.rule_id} 整改建议**: `{', '.join(r.recommendation_ids)}`")
            if r.verification_ids:
                lines.append(f"- **{r.rule_id} 核查建议**: `{', '.join(r.verification_ids)}`")
        for obs in (o for o in report.observations if o.domain == domain):
            lines.append(f"- **观察 {obs.observation_id}**: {obs.text} 引用 `{obs.fact_ref}`。")
        lines.append("")
    lines += ["## 无法判断与采集限制", ""] + [f"- {item}" for item in report.limitations]
    lines += ["", "## 附录", "", "| 共享规则 | subject | 结论 | reason_code |", "| --- | --- | --- | --- |"]
    lines += [f"| {r.rule_id} | `{r.subject}` | {LABELS[r.conclusion]} | `{r.reason_code}` |" for r in report.shared_result_references]
    lines += [""] + [f"- {item}" for item in report.appendix]
    return "\n".join(lines).rstrip() + "\n"


@dataclass
class Block:
    kind: str
    text: str = ""
    level: int = 0
    items: list[str] = field(default_factory=list)
    ordered: bool = False
    rows: list[list[str]] = field(default_factory=list)


class MarkdownValidationError(Exception):
    pass


def safe_error(line: int, construct: str) -> MarkdownValidationError:
    return MarkdownValidationError(f"line {line}: unsupported or invalid {construct}")


def validate_inline(line_no: int, text: str) -> None:
    if any((ord(c) < 32 and c not in "\t\n\r") or ord(c) == 127 for c in text):
        raise safe_error(line_no, "raw control character")
    if text.count("`") % 2:
        raise safe_error(line_no, "unclosed inline code")
    remainder = re.sub(r"`[^`\n]+`", "", text)
    if re.search(r"<[^>]+>|<!--", remainder): raise safe_error(line_no, "HTML")
    if re.search(r"!\[", remainder): raise safe_error(line_no, "image")
    for match in list(re.finditer(r"\[([^\]]+)\]\(([^)]+)\)", remainder))[::-1]:
        parsed = urlparse(match.group(2))
        if parsed.scheme not in {"http", "https"} or not parsed.netloc: raise safe_error(line_no, "unsafe link")
        remainder = remainder[:match.start()] + remainder[match.end():]
    if "[" in remainder or "]" in remainder: raise safe_error(line_no, "malformed link")
    remainder = re.sub(r"\*\*[^*\n]+\*\*", "", remainder)
    if "**" in remainder: raise safe_error(line_no, "unclosed bold")
    if re.search(r"(?<!\*)\*(?!\*)", remainder) or re.search(r"(?<![\w_])_[^_\n]+_(?![\w_])", remainder): raise safe_error(line_no, "emphasis")


def parse_markdown(markdown: str) -> list[Block]:
    lines, blocks, paragraph, items, table = markdown.splitlines(), [], [], [], []
    ordered = False
    table_start = 0
    last_heading = 0
    def flush() -> None:
        nonlocal paragraph, items, table
        if paragraph: blocks.append(Block("paragraph", text=" ".join(paragraph))); paragraph = []
        if items: blocks.append(Block("list", items=items, ordered=ordered)); items = []
        if table:
            width = len(table[0])
            if len(table) < 3 or width < 2 or any(len(row) != width for row in table): raise safe_error(table_start, "ragged table")
            if not all(re.fullmatch(r":?-{3,}:?", c) for c in table[1]): raise safe_error(table_start + 1, "table separator")
            blocks.append(Block("table", rows=[table[0], *table[2:]])); table = []
    for n, raw in enumerate(lines, 1):
        line = raw.rstrip()
        validate_inline(n, line)
        if re.fullmatch(r" {0,3}([-*_])(?:\s*\1){2,}\s*", line): raise safe_error(n, "horizontal rule")
        if line.startswith(("```", "~~~")): raise safe_error(n, "fenced code block")
        if line.startswith(">"): raise safe_error(n, "blockquote")
        if re.match(r"^( {4}|\t)\S", line): raise safe_error(n, "indented code block")
        if not line.strip(): flush(); continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            flush(); level = len(heading.group(1))
            if level not in {1, 2, 3} or (last_heading and level > last_heading + 1): raise safe_error(n, "heading level")
            last_heading = level; blocks.append(Block("heading", text=heading.group(2), level=level)); continue
        if re.match(r"^\s+([-*+]|\d+[.)])\s+", line): raise safe_error(n, "nested list")
        bullet, numbered = re.match(r"^-\s+(.+)$", line), re.match(r"^\d+\.\s+(.+)$", line)
        if bullet or numbered:
            if paragraph or table: flush()
            now_ordered = bool(numbered)
            if items and now_ordered != ordered: flush()
            ordered = now_ordered; items.append((numbered or bullet).group(1)); continue
        if line.startswith("|") or line.endswith("|"):
            if paragraph or items: flush()
            table_start = table_start or n
            table.append([c.strip() for c in line.strip().strip("|").split("|")]); continue
        if table: flush()
        paragraph.append(line.strip())
    flush()
    return blocks


def plain_inline(text: str) -> str:
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return re.sub(r"\[([^\]]+)\]\((https?://[^)]+)\)", r"\1 (\2)", text)


def add_inline(paragraph, text: str) -> None:
    token_re = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\(https?://[^)]+\))")
    pos = 0
    for match in token_re.finditer(text):
        paragraph.add_run(text[pos:match.start()]); token = match.group(0)
        if token.startswith("**"): run = paragraph.add_run(token[2:-2]); run.bold = True
        elif token.startswith("`"): run = paragraph.add_run(token[1:-1]); run.font.name = "Courier New"
        else: run = paragraph.add_run(plain_inline(token)); run.underline = True
        pos = match.end()
    paragraph.add_run(text[pos:])


def markdown_to_docx(markdown: str, out: Path, revised_export: bool) -> None:
    blocks = parse_markdown(markdown)  # Validate completely before creating/writing output.
    document = Document()
    if revised_export:
        run = document.add_paragraph().add_run("用户修订版：DOCX 内容来源为用户修订 Markdown。")
        run.bold = True
    for block in blocks:
        if block.kind == "heading": document.add_heading(block.text, level=block.level)
        elif block.kind == "paragraph": add_inline(document.add_paragraph(), block.text)
        elif block.kind == "list":
            for item in block.items: add_inline(document.add_paragraph(style="List Number" if block.ordered else "List Bullet"), item)
        else:
            table = document.add_table(rows=1, cols=len(block.rows[0])); table.style = "Table Grid"
            for i, text in enumerate(block.rows[0]): run = table.rows[0].cells[i].paragraphs[0].add_run(plain_inline(text)); run.bold = True
            for row in block.rows[1:]:
                cells = table.add_row().cells
                for i, text in enumerate(row): add_inline(cells[i].paragraphs[0], text)
    out.parent.mkdir(parents=True, exist_ok=True)
    temp = out.with_suffix(".docx.tmp"); document.save(temp); temp.replace(out)


def normalized_markdown(markdown: str) -> list[tuple]:
    result = []
    for b in parse_markdown(markdown):
        if b.kind == "heading": result.append(("heading", b.level, b.text))
        elif b.kind == "paragraph": result.append(("paragraph", plain_inline(b.text)))
        elif b.kind == "list": result.append(("list", b.ordered, tuple(plain_inline(x) for x in b.items)))
        else: result.append(("table", tuple(tuple(plain_inline(c) for c in row) for row in b.rows)))
    return result


def normalized_docx(path: Path) -> list[tuple]:
    document, result, pending, pending_ordered = Document(path), [], [], None
    def flush_list() -> None:
        nonlocal pending
        if pending: result.append(("list", pending_ordered, tuple(pending))); pending = []
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = next(p for p in document.paragraphs if p._p is child)
            text, style = paragraph.text.strip(), paragraph.style.name
            if not text or text.startswith("用户修订版"): continue
            if style in {"List Bullet", "List Number"}:
                current = style == "List Number"
                if pending and current != pending_ordered: flush_list()
                pending_ordered = current; pending.append(text); continue
            flush_list()
            if style.startswith("Heading"): result.append(("heading", int(style.split()[-1]), text))
            else: result.append(("paragraph", text))
        elif child.tag == qn("w:tbl"):
            flush_list(); table = next(t for t in document.tables if t._tbl is child)
            result.append(("table", tuple(tuple(c.text.strip() for c in row.cells) for row in table.rows)))
    flush_list(); return result


def write_artifacts(name: str, report: ReportView, revised: bool = False) -> tuple[Path, Path, str]:
    md = render_report_view_markdown(report)
    if revised:
        md = md.replace("**报告来源**: 系统生成报告", "**报告来源**: 用户修订报告")
        md = md.replace("| user_revised | false |", "| user_revised | true |")
        md = md.replace("观察项为非判断性事实记录，不计入规则结果数量。", "观察项为非判断性事实记录，不计入规则结果数量。\n\n用户补充说明：已复核变更冻结期。")
        md = md.replace("- 五类结论分别计数；观察项不进入统计。", "- 五类结论分别计数；观察项不进入统计。\n- 用户新增列表项：已复核章节顺序。")
        md = md.replace("| 共享规则 | subject | 结论 | reason_code |", "| 共享规则 | subject | 结论 | reason_code | 复核 |")
        md = md.replace("| --- | --- | --- | --- |\n| INSTANCE.DISCOVERY", "| --- | --- | --- | --- | --- |\n| INSTANCE.DISCOVERY")
        for shared in ("INSTANCE.DISCOVERY", "HOST.VISIBILITY", "HOST.MEMORY_CAPACITY"):
            md = re.sub(rf"^(\| {re.escape(shared)} \| .*? \| `[^`]+`)( \|)$", r"\1 | 用户保留 |", md, flags=re.MULTILINE)
    md_path, docx_path = TMP / f"{name}.md", TMP / f"{name}.docx"
    md_path.write_text(md, encoding="utf-8"); markdown_to_docx(md, docx_path, revised)
    return md_path, docx_path, md


def all_text(path: Path) -> str:
    doc = Document(path)
    return "\n".join([p.text for p in doc.paragraphs] + [c.text for t in doc.tables for r in t.rows for c in r.cells])


def check(condition: bool, message: str) -> None:
    if not condition: raise AssertionError(message)


def scenario_canonical_normal_9_and_10() -> None:
    for instance, version in (("A", "9.0.89"), ("B", "10.1.24")):
        report = make_report(instance, version); _, docx, md = write_artifacts(f"normal-{instance}", report)
        check(all(r.rule_id in {"HOST.FILESYSTEM_CAPACITY", "SECURITY.VERSION_LIFECYCLE", "JVM.VERSION_COMPATIBILITY", "CONNECTOR.TIMEOUTS", "SECURITY.SHUTDOWN_PORT", "SECURITY.AJP", "APPLICATION.CONTEXT_CONFLICT", "LOGGING.ROTATION"} for r in report.rules), "noncanonical rule")
        check("警告" in md and "预警" not in md, "wrong conclusion vocabulary")
        check(version in all_text(docx), "version missing")
        check(report.summary_counts["not_applicable"] == 1, "true no-AJP not_applicable missing")
        check("CPU 核数" in md and "应用数量" in md and "日志文件大小" in md, "observations missing")


def scenario_canonical_degraded() -> None:
    report = make_report("B", "10.1.24", True); _, docx, md = write_artifacts("degraded", report)
    check(all_text(docx).find("警告") >= 0 and all_text(docx).find("异常") >= 0 and all_text(docx).find("无法判断") >= 0, "degraded conclusions missing")
    for r in report.rules: r.validate()
    check("核查建议" in md and "整改建议" in md, "recommendation/verification separation absent")
    check("SR-JVM" not in md and "SR-CONN" not in md, "instance facts mislabeled shared")


def scenario_independent_reports() -> None:
    _, a, _ = write_artifacts("independent-A", make_report("A", "9.0.89")); _, b, _ = write_artifacts("independent-B", make_report("B", "10.1.24", True))
    check("tomcat-B" not in all_text(a) and "18080" not in all_text(a), "B bled into A")
    check("tomcat-A" not in all_text(b), "A bled into B")


def scenario_user_revision_and_semantics() -> None:
    report = make_report("A", "9.0.89"); check(not report.user_revised, "initial view revision flag")
    _, docx, md = write_artifacts("user-revised", report, True)
    check("用户修订报告" in md and "user_revised | true" in md and "用户修订版" in all_text(docx), "revised markers missing")
    check(normalized_markdown(md) == normalized_docx(docx), "ordered Markdown/DOCX semantics differ")


def scenario_invalid_markdown() -> None:
    cases = {
        "html": "# A\n<div>x</div>", "image": "# A\n![x](https://e.test/x)", "quote": "# A\n> x", "fence": "# A\n```\nx\n```", "indent": "# A\n    x",
        "nested": "# A\n- x\n  - y", "ragged": "# A\n| A | B |\n| --- | --- |\n| 1 |", "heading": "#### A", "unsafe": "# A\n[x](javascript:x)", "control": "# A\nx\x01",
        "bold_open": "# A\n**open", "bold_close": "# A\nclose**", "code_open": "# A\n`open", "link_open": "# A\n[x](https://e.test", "emphasis_star": "# A\n*x*", "emphasis_under": "# A\n_x_", "horizontal": "# A\n---",
    }
    for name, markdown in cases.items():
        out = TMP / f"invalid-{name}.docx"
        try: markdown_to_docx(markdown, out, False)
        except MarkdownValidationError as exc:
            check(not out.exists(), f"invalid output written: {name}")
            check("javascript" not in str(exc) and "<div" not in str(exc), "unsafe content echoed")
        else: raise AssertionError(f"invalid accepted: {name}")


def scenario_inline_literal_code() -> None:
    markdown = "# A\n\n字面量 `*not emphasis*`、`<tag>` 与 `[x](javascript:x)` 必须保真。\n"
    out = TMP / "inline-literal-code.docx"
    markdown_to_docx(markdown, out, False)
    text = all_text(out)
    check("*not emphasis*" in text and "<tag>" in text and "[x](javascript:x)" in text, "inline code literal content lost")


def scenario_docx_structure_order() -> None:
    _, path, md = write_artifacts("structure", make_report("A", "9.0.89"), True)
    blocks = normalized_docx(path)
    headings = [b[2] for b in blocks if b[0] == "heading"]
    check([h for h in headings if h in MAJOR_SECTIONS] == list(MAJOR_SECTIONS), "major heading order")
    overview = blocks.index(("heading", 2, "实例概况")); summary = blocks.index(("heading", 2, "结论摘要"))
    check(blocks[overview + 1][0] == "table" and blocks[summary + 1][0] == "table", "tables not immediately after intended headings")
    with zipfile.ZipFile(path) as archive: xml = archive.read("word/document.xml").decode("utf-8")
    overview_pos, metadata_pos, summary_pos = xml.find("实例概况"), xml.find("protocol_version"), xml.find("结论摘要")
    status_pos = xml.find("状态", summary_pos)
    check(overview_pos < metadata_pos < summary_pos < status_pos, "XML structural order wrong")


def main() -> int:
    if TMP.exists(): shutil.rmtree(TMP)
    TMP.mkdir(parents=True)
    scenarios = [scenario_canonical_normal_9_and_10, scenario_canonical_degraded, scenario_independent_reports, scenario_user_revision_and_semantics, scenario_invalid_markdown, scenario_inline_literal_code, scenario_docx_structure_order]
    results = []
    for scenario in scenarios:
        try: scenario(); results.append((scenario.__name__, "PASS", ""))
        except Exception as exc: results.append((scenario.__name__, "FAIL", str(exc)))
    for name, status, detail in results: print(f"{status} {name}" + (f" - {detail}" if detail else ""))
    passed = sum(status == "PASS" for _, status, _ in results)
    print(f"SUMMARY {passed}/{len(results)} scenarios passed"); print(f"OUTPUT {TMP}")
    return 0 if passed == len(results) else 1


if __name__ == "__main__": raise SystemExit(main())
